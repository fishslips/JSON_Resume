import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load JSON data
with open('resume_jsons/tim_severance_resume.json', 'r') as file:
    resume_data = json.load(file)

# Create a new Document
doc = Document()
# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Add name
name_heading = doc.add_heading(resume_data['name'], level=1)
name_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add contact information
contact_info = resume_data['contact']
contact_line = f"{contact_info['email']} | {contact_info['phone']} | {contact_info['linkedin']}"
contact_paragraph = doc.add_paragraph(contact_line)
contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add summary
doc.add_heading('Summary', level=2)
doc.add_paragraph(resume_data['summary'])

# Add education
doc.add_heading('Education', level=2)
for edu in resume_data['education']:
    doc.add_paragraph(f"{edu['degree']} - {edu['institution']} ({edu['year']})")

# Add experience
doc.add_heading('Experience', level=2)
for exp in resume_data['experience']:
    doc.add_heading(exp['job_title'], level=3)
    doc.add_paragraph(f"{exp['company']} ({exp['start_date']} - {exp['end_date']})")
    for responsibility in exp['responsibilities']:
        doc.add_paragraph(f"- {responsibility}")

# Add skills
doc.add_heading('Skills', level=2)
for skill_category, skills in resume_data['skills'].items():
    doc.add_heading(skill_category, level=3)
    doc.add_paragraph(", ".join(skills))

# Add projects
doc.add_heading('Projects', level=2)
for project in resume_data['projects']:
    doc.add_heading(project['name'], level=3)
    doc.add_paragraph(project['description'])
    doc.add_paragraph(f"Technologies: {', '.join(project['technologies'])}")
    doc.add_paragraph(f"Duration: {project['start_date']} - {project['end_date']}")

# Add publications
doc.add_heading('Publications', level=2)
for pub in resume_data['publications']:
    doc.add_heading(pub['title'], level=3)
    doc.add_paragraph(f"Authors: {', '.join(pub['authors'])}")
    doc.add_paragraph(f"Publication: {pub['publication']} ({pub['year']})")

# Add certifications
doc.add_heading('Certifications', level=2)
for cert in resume_data['certifications']:
    doc.add_paragraph(f"{cert['name']} - {cert['institution']} ({cert['year']})")

# Save the document
doc.save('resume_docx/Tim_Severance_Resume.docx')